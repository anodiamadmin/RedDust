# services/rag/parsers.py — File parser layer
#
# Parses .xlsx and .docx files into semantically meaningful chunks with metadata.
#
# .docx files in this project do NOT use Word heading styles — all paragraphs
# are styled 'Normal'. Heading hierarchy is inferred differently per document:
#
#   TraditionalWisdom.docx:
#       bold + size 24  → H1 (culture entry — one chunk per culture)
#       bold + size 18  → H2 (sub-section, accumulates into culture chunk)
#       Preamble before first H1 discarded via cultures_started flag.
#       H2 "The RedDust lesson" → metadata['reddust_relevance']
#
#   ScientificResearchKnowledgeBase.docx (and any future research docs):
#       No detectable formatting signals. Heading detection uses a generalised
#       heuristic — short line, no sentence-ending punctuation, not a bullet/
#       citation. H1 additionally requires year-in-parentheses or named title,
#       AND must not start with a relational prefix ("Relationship to",
#       "Integration with") which are always H2 sub-sections.
#       H2 "Relevance to RedDust" → metadata['reddust_relevance']
#
#   SoulScoreMethodology.docx:
#       Domain: soul_score_context. Do NOT ingest until Steps 18-22.

from pathlib import Path
import re
from typing import TypedDict

import openpyxl
from docx import Document


# ---------------------------------------------------------------------------
# Shared types
# ---------------------------------------------------------------------------

class Chunk(TypedDict):
    text: str
    metadata: dict  # filename, source; optionally: culture, reddust_relevance


# ---------------------------------------------------------------------------
# Shared constants
# ---------------------------------------------------------------------------

_RELEVANCE_MARKERS = {
    "relevance to reddust",
    "the reddust lesson",
    "reddust lesson",
    "reddust relevance",
}

# Regex: 4-digit year in parentheses — present in research entry titles
_YEAR_RE = re.compile(r'\(\d{4}\)')

# Punctuation that signals a body sentence (not a heading)
_SENTENCE_ENDINGS = {'.', ',', ';', ':'}

# Prefixes that signal a bullet/list item (not a heading)
_BULLET_PREFIXES = ('•', '-', '*', '–', '—')

# Substrings that signal a citation line (not a heading)
_CITATION_SIGNALS = ('doi:', 'http', 'pp.', 'vol.', 'doi.org', 'org/')

# Prefixes that mark a relational H2 sub-section — never an H1 entry title.
# These lines may contain a year in parentheses but are always sub-sections.
_RELATIONAL_PREFIXES = (
    'relationship to',
    'integration with',
    'related to',
    'connection to',
)

# Known named H1 titles that have no year in parentheses.
# Add here if a future doc introduces a model/framework title without a year.
_NAMED_H1_TITLES = {
    "juslin's brecvema model",
}


# ---------------------------------------------------------------------------
# Generalised heading heuristic
# ---------------------------------------------------------------------------

def _is_heading_line(text: str) -> bool:
    """
    Returns True if a paragraph looks like a heading rather than body text.
    Criteria (ALL must be true):
      1. Length <= 80 characters
      2. Does not end with sentence-ending punctuation (. , ; :)
      3. Not a bullet or numbered list item
      4. Not a citation line
      5. Not a star-rating line (★)
    """
    stripped = text.strip()

    if len(stripped) > 80:
        return False
    if stripped and stripped[-1] in _SENTENCE_ENDINGS:
        return False
    if stripped.startswith(_BULLET_PREFIXES):
        return False
    if re.match(r'^\d+\.', stripped):
        return False
    lower = stripped.lower()
    if any(sig in lower for sig in _CITATION_SIGNALS):
        return False
    if stripped.startswith('★'):
        return False

    return True


def _is_scientific_h1(text: str) -> bool:
    """
    True if paragraph is a research entry title (H1 chunk boundary).
    Must pass _is_heading_line AND:
      - contain a year in parentheses OR match a named model title
      AND
      - NOT start with a relational prefix (those are always H2 sub-sections)
    """
    if not _is_heading_line(text):
        return False

    stripped = text.strip()
    lower = stripped.lower()

    # Relational prefixes are always H2 — block them from becoming H1
    if any(lower.startswith(prefix) for prefix in _RELATIONAL_PREFIXES):
        return False

    if _YEAR_RE.search(stripped):
        return True
    if lower in _NAMED_H1_TITLES:
        return True

    return False


def _is_relevance_heading(text: str) -> bool:
    """True if heading marks a 'Relevance to RedDust' / 'RedDust lesson' section."""
    return text.strip().lower() in _RELEVANCE_MARKERS


# ---------------------------------------------------------------------------
# Formatting helpers (TraditionalWisdom parser)
# ---------------------------------------------------------------------------

def _para_formatting(para) -> tuple[bool, int | None]:
    """Return (is_bold, font_size_pt) for a paragraph."""
    bold = any(run.bold for run in para.runs) if para.runs else False
    font_size = next((run.font.size for run in para.runs if run.font.size), None)
    size_pt = round(font_size / 12700) if font_size else None
    return bold, size_pt


# ---------------------------------------------------------------------------
# .xlsx parser
# ---------------------------------------------------------------------------

def parse_xlsx(path: str | Path) -> list[Chunk]:
    """
    One chunk per non-empty row. Column values joined with ' | '.
    Metadata: filename, source (sheet name).
    """
    path = Path(path)
    wb = openpyxl.load_workbook(path, data_only=True)
    chunks: list[Chunk] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows(values_only=True):
            text = " | ".join(str(cell) for cell in row if cell is not None)
            if text.strip():
                chunks.append(Chunk(
                    text=text.strip(),
                    metadata={"filename": path.name, "source": sheet_name},
                ))

    return chunks


# ---------------------------------------------------------------------------
# TraditionalWisdom parser — formatting-aware (bold + font size)
# ---------------------------------------------------------------------------

def _parse_docx_by_formatting(
    path: Path,
    h1_size: int,
    h2_size: int,
    culture_mode: bool,
) -> list[Chunk]:
    """
    Infers heading level from bold+size. Used for TraditionalWisdom.docx.
    Preamble (before first H1) discarded via cultures_started flag.
    """
    doc = Document(str(path))
    chunks: list[Chunk] = []

    current_heading: str = "root"
    current_culture: str | None = None
    current_body: list[str] = []
    current_relevance: list[str] = []
    in_relevance_section: bool = False
    cultures_started: bool = False

    def _flush() -> None:
        if not cultures_started:
            return
        body = "\n".join(current_body).strip()
        relevance = "\n".join(current_relevance).strip()
        if not body and not relevance:
            return
        meta: dict = {"filename": path.name, "source": current_heading}
        if current_culture:
            meta["culture"] = current_culture
        if relevance:
            meta["reddust_relevance"] = relevance
        chunks.append(Chunk(text=body or relevance, metadata=meta))

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        bold, size_pt = _para_formatting(para)

        if bold and size_pt == h1_size:
            _flush()
            cultures_started = True
            current_heading = text
            current_culture = text if culture_mode else None
            current_body = []
            current_relevance = []
            in_relevance_section = False
            continue

        if not cultures_started:
            continue

        if bold and size_pt == h2_size:
            if _is_relevance_heading(text):
                in_relevance_section = True
            else:
                in_relevance_section = False
                if culture_mode:
                    current_body.append(f"\n## {text}")
                else:
                    _flush()
                    current_heading = text
                    current_body = []
                    current_relevance = []
            continue

        if in_relevance_section:
            current_relevance.append(text)
        else:
            current_body.append(text)

    _flush()
    return chunks


# ---------------------------------------------------------------------------
# Scientific research parser — generalised heuristic heading detection
# ---------------------------------------------------------------------------

def _parse_scientific_docx(path: Path) -> list[Chunk]:
    """
    Parses research docs where all paragraphs share Normal style with no
    detectable formatting differences.

    H1 = _is_scientific_h1() — year in parentheses or named title, not relational
    H2 = _is_heading_line() and not H1 — accumulates into current chunk body
         except "Relevance to RedDust" which goes to relevance bucket

    One chunk per research entry (H1). Sub-sections merge into it for full
    study context at retrieval time.
    """
    doc = Document(str(path))
    chunks: list[Chunk] = []

    current_heading: str = "root"
    current_body: list[str] = []
    current_relevance: list[str] = []
    in_relevance_section: bool = False
    entries_started: bool = False

    def _flush() -> None:
        if not entries_started:
            return
        body = "\n".join(current_body).strip()
        relevance = "\n".join(current_relevance).strip()
        if not body and not relevance:
            return
        meta: dict = {"filename": path.name, "source": current_heading}
        if relevance:
            meta["reddust_relevance"] = relevance
        chunks.append(Chunk(text=body or relevance, metadata=meta))

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if _is_scientific_h1(text):
            _flush()
            entries_started = True
            current_heading = text
            current_body = []
            current_relevance = []
            in_relevance_section = False
            continue

        if not entries_started:
            continue

        if _is_heading_line(text):
            if _is_relevance_heading(text):
                in_relevance_section = True
            else:
                in_relevance_section = False
                current_body.append(f"\n## {text}")
            continue

        if in_relevance_section:
            current_relevance.append(text)
        else:
            current_body.append(text)

    _flush()
    return chunks


# ---------------------------------------------------------------------------
# Named parsers
# ---------------------------------------------------------------------------

def parse_traditional_wisdom(path: str | Path) -> list[Chunk]:
    """
    TraditionalWisdom.docx
    bold+24=culture H1, bold+18=sub-section H2.
    'The RedDust lesson' → metadata['reddust_relevance'].
    Domain at ingestion: music_research
    """
    return _parse_docx_by_formatting(Path(path), h1_size=24, h2_size=18, culture_mode=True)


def parse_scientific_research(path: str | Path) -> list[Chunk]:
    """
    ScientificResearchKnowledgeBase.docx and any future flat-formatted research doc.
    Generalised heuristic heading detection — no hardcoded sub-section names.
    One chunk per research entry. 'Relevance to RedDust' → metadata['reddust_relevance'].
    Domain at ingestion: music_research

    Adding a new research doc: create a wrapper calling _parse_scientific_docx(Path(path)).
    If the doc has a named model title with no year, add it to _NAMED_H1_TITLES.
    If it has new relational H2 prefixes, add them to _RELATIONAL_PREFIXES.
    """
    return _parse_scientific_docx(Path(path))


def parse_soul_score_context(path: str | Path) -> list[Chunk]:
    """
    SoulScoreMethodology.docx — narrative + scientific backing.
    Do NOT ingest until Steps 18-22 (wellbeing tracking) are built.
    Domain at ingestion: soul_score_context
    """
    return _parse_docx_by_formatting(Path(path), h1_size=24, h2_size=18, culture_mode=False)
